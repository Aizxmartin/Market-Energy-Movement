
import argparse
import os
from datetime import datetime, timedelta
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REQUIRED_COLUMNS = [
    "Mls Status",
    "Close Date",
    "List Price",
    "Close Price",
    # "Seller Concessions",  # now optional
    # "DaysInMLS",          # allow alias "Days in MLS"
]

ACTIVE_KEYS = ["active", "coming soon", "back on market", "a/"]
PENDING_KEYS = ["pending", "under contract", "a/i", "accepting backup"]
SOLD_KEYS = ["closed", "sold"]

def to_num(x):
    if pd.isna(x):
        return np.nan
    s = str(x).replace("$", "").replace(",", "").strip()
    try:
        return float(s)
    except Exception:
        return np.nan

def bucket_status(s: str) -> str:
    if pd.isna(s):
        return "Unknown"
    s = str(s).strip().lower()
    # Active if contains any active key AND not pending/UC
    if any(k in s for k in ACTIVE_KEYS) and not any(k in s for k in ["under contract", "pending"]):
        return "Active"
    if any(k in s for k in PENDING_KEYS):
        return "Pending"
    if any(k in s for k in SOLD_KEYS):
        return "Sold"
    if s == "active":
        return "Active"
    return "Other"

def get_dom_series(df: pd.DataFrame) -> pd.Series:
    dom_col = None
    if "DaysInMLS" in df.columns:
        dom_col = "DaysInMLS"
    elif "Days in MLS" in df.columns:
        dom_col = "Days in MLS"
    if dom_col is None:
        return pd.Series([np.nan] * len(df))
    try:
        return pd.to_numeric(df[dom_col].astype(str).str.replace(",", "", regex=False), errors="coerce")
    except Exception:
        return pd.to_numeric(df[dom_col], errors="coerce")

def main():
    parser = argparse.ArgumentParser(description="Generate Momentum Report from MLS CSV (fixed field names where possible).")
    parser.add_argument("csv_path", help="Path to the CSV file with required fields.")
    parser.add_argument("--days", type=int, default=90, help="Window (days) for solds, default 90.")
    args = parser.parse_args()

    csv_path = args.csv_path
    if not os.path.exists(csv_path):
        raise SystemExit(f"File not found: {csv_path}")

    df = pd.read_csv(csv_path, dtype=str)

    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        raise SystemExit(f"Missing required columns: {missing}")

    # Optional columns setup
    has_concessions = "Seller Concessions" in df.columns

    # Prep
    df["_bucket"] = df["Mls Status"].apply(bucket_status)
    df["_close_date"] = pd.to_datetime(df["Close Date"], errors="coerce", infer_datetime_format=True)
    df["_list_price"] = df["List Price"].apply(to_num)
    df["_close_price"] = df["Close Price"].apply(to_num)
    df["_concessions"] = df["Seller Concessions"].apply(to_num) if has_concessions else 0.0
    df["_dom"] = get_dom_series(df)

    # Solds window
    today = datetime.today()
    window_start = today - timedelta(days=args.days)
    is_sold_window = (df["_bucket"] == "Sold") & (df["_close_date"] >= window_start)

    # Counts
    active_count = int((df["_bucket"] == "Active").sum())
    pending_count = int((df["_bucket"] == "Pending").sum())
    sold_window_count = int(is_sold_window.sum())

    # MOI
    den = (sold_window_count / 3.0) + pending_count
    moi = (active_count / den) if den > 0 else np.nan

    # Ranges
    active_prices = df.loc[df["_bucket"] == "Active", "_list_price"].dropna()
    pending_prices = df.loc[df["_bucket"] == "Pending", "_list_price"].dropna()

    net_price = df["_close_price"] - (df["_concessions"] if has_concessions else 0)
    sold_net_prices = net_price.loc[is_sold_window].dropna()

    active_min = float(active_prices.min()) if not active_prices.empty else np.nan
    active_max = float(active_prices.max()) if not active_prices.empty else np.nan
    pending_min = float(pending_prices.min()) if not pending_prices.empty else np.nan
    pending_max = float(pending_prices.max()) if not pending_prices.empty else np.nan
    sold_net_min = float(sold_net_prices.min()) if not sold_net_prices.empty else np.nan
    sold_net_max = float(sold_net_prices.max()) if not sold_net_prices.empty else np.nan

    # DaysInMLS stats for solds in window
    dom_series = df.loc[is_sold_window, "_dom"].dropna()
    avg_dom = float(dom_series.mean()) if not dom_series.empty else np.nan
    median_dom = float(dom_series.median()) if not dom_series.empty else np.nan

    # Chart
    out_dir = os.path.dirname(os.path.abspath(csv_path)) or "."
    date_stamp = today.strftime("%Y%m%d")
    chart_path = os.path.join(out_dir, f"momentum_counts_{date_stamp}.png")
    summary_counts = {"Active": active_count, "Pending": pending_count, f"Solds ({args.days}d)": sold_window_count}

    # Create bar chart
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.bar(list(summary_counts.keys()), list(summary_counts.values()))
    ax.set_title("Status Counts")
    ax.set_ylabel("Count")
    plt.tight_layout()
    fig.savefig(chart_path, dpi=150)
    plt.close(fig)

    # DOCX report
    doc = Document()
    title = doc.add_paragraph("Momentum Report")
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pdate = doc.add_paragraph(f"Report Date: {today.strftime('%B %d, %Y')}")
    pdate.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph("Summary", style=None).runs[0].bold = True
    doc.add_paragraph(f"Active: {active_count}")
    doc.add_paragraph(f"Pending: {pending_count}")
    doc.add_paragraph(f"Solds (last {args.days}d): {sold_window_count}")
    if not np.isnan(moi):
        doc.add_paragraph(f"Months of Inventory (MOI): {moi:.3f}")
    else:
        doc.add_paragraph("Months of Inventory (MOI): N/A")

    doc.add_paragraph("Price Ranges", style=None).runs[0].bold = True
    doc.add_paragraph(
        f"Active List Price Range: ${active_min:,.0f} – ${active_max:,.0f}"
        if not (np.isnan(active_min) or np.isnan(active_max)) else "Active List Price Range: N/A"
    )
    doc.add_paragraph(
        f"Pending List Price Range: ${pending_min:,.0f} – ${pending_max:,.0f}"
        if not (np.isnan(pending_min) or np.isnan(pending_max)) else "Pending List Price Range: N/A"
    )
    doc.add_paragraph(
        f"Closed Net Price Range ({args.days}d): ${sold_net_min:,.0f} – ${sold_net_max:,.0f}"
        if not (np.isnan(sold_net_min) or np.isnan(sold_net_max)) else f"Closed Net Price Range ({args.days}d): N/A"
    )

    doc.add_paragraph("Days in MLS (Closed, Window)", style=None).runs[0].bold = True
    doc.add_paragraph(f"Average DaysInMLS: {avg_dom:.1f}" if not np.isnan(avg_dom) else "Average DaysInMLS: N/A")
    doc.add_paragraph(f"Median DaysInMLS: {median_dom:.1f}" if not np.isnan(median_dom) else "Median DaysInMLS: N/A")

    doc.add_paragraph("")
    doc.add_paragraph("Status Counts Chart", style=None).runs[0].bold = True
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))

    doc.add_paragraph("")
    doc.add_paragraph("Formula", style=None).runs[0].bold = True
    doc.add_paragraph(f"MOI = Active / ((Solds_{args.days}d / 3) + Pending)")

    doc.add_paragraph("")
    doc.add_paragraph("Disclaimer", style=None).runs[0].bold = True
    doc.add_paragraph(
        "This report is generated from the provided MLS export. Field names are expected to match the fixed set; "
        "optional fields like Seller Concessions and DaysInMLS may be omitted."
    )

    out_docx = os.path.join(out_dir, f"Momentum_Report_{date_stamp}.docx")
    doc.save(out_docx)

    print("Report saved:", out_docx)
    print("Chart saved:", chart_path)

if __name__ == "__main__":
    main()
