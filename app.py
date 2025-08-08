import io
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

st.set_page_config(page_title="Momentum Report", page_icon="📊", layout="centered")

st.title("📊 Momentum Report (CSV ➜ DOCX)")

st.markdown("""
Upload your MLS CSV and generate a **Momentum Report** as a Word (.docx) file.

**Required columns (exact names):**
- `Mls Status`, `Close Date`, `List Price`, `Close Price`

**Optional columns:**
- `Seller Concessions` (defaults to 0 if missing)
- `DaysInMLS` or `Days in MLS` (for Avg/Median Days on MLS in closed window)
""")

uploaded = st.file_uploader("Upload CSV", type=["csv"])
days_window = st.number_input("Closed sales lookback (days)", min_value=30, max_value=365, value=90, step=5)

def to_num(x):
    if pd.isna(x):
        return np.nan
    s = str(x).replace("$", "").replace(",", "").strip()
    try:
        return float(s)
    except Exception:
        return np.nan

ACTIVE_KEYS = ["active", "coming soon", "back on market", "a/"]
PENDING_KEYS = ["pending", "under contract", "a/i", "accepting backup"]
SOLD_KEYS = ["closed", "sold"]

def bucket_status(s: str) -> str:
    if pd.isna(s):
        return "Unknown"
    s = str(s).strip().lower()
    # Active if contains any active key AND not pending/UC
    if any(k in s for k in ACTIVE_KEYS) and not any(k in s for k in ["under contract", "pending"]):
        return "Active"
    if any(k in s for k in PENDING_KEYS):
        return "Pending"
    if any(k in s for k in SOLD_KEYS):
        return "Sold"
    if s == "active":
        return "Active"
    return "Other"

def get_dom_series(df: pd.DataFrame) -> pd.Series:
    dom_col = None
    if "DaysInMLS" in df.columns:
        dom_col = "DaysInMLS"
    elif "Days in MLS" in df.columns:
        dom_col = "Days in MLS"
    if dom_col is None:
        return pd.Series([np.nan] * len(df))
    try:
        return pd.to_numeric(df[dom_col].astype(str).str.replace(",", "", regex=False), errors="coerce")
    except Exception:
        return pd.to_numeric(df[dom_col], errors="coerce")

def build_docx(df: pd.DataFrame, days: int) -> bytes:
    has_concessions = "Seller Concessions" in df.columns

    df = df.copy()
    df["_bucket"] = df["Mls Status"].apply(bucket_status)
    df["_close_date"] = pd.to_datetime(df["Close Date"], errors="coerce", infer_datetime_format=True)
    df["_list_price"] = df["List Price"].apply(to_num)
    df["_close_price"] = df["Close Price"].apply(to_num)
    df["_concessions"] = df["Seller Concessions"].apply(to_num) if has_concessions else 0.0
    df["_dom"] = get_dom_series(df)

    today = datetime.today()
    window_start = today - timedelta(days=days)
    is_sold_window = (df["_bucket"] == "Sold") & (df["_close_date"] >= window_start)

    active_count = int((df["_bucket"] == "Active").sum())
    pending_count = int((df["_bucket"] == "Pending").sum())
    sold_window_count = int(is_sold_window.sum())

    den = (sold_window_count / 3.0) + pending_count
    moi = (active_count / den) if den > 0 else np.nan

    active_prices = df.loc[df["_bucket"] == "Active", "_list_price"].dropna()
    pending_prices = df.loc[df["_bucket"] == "Pending", "_list_price"].dropna()
    net_price = df["_close_price"] - (df["_concessions"] if has_concessions else 0)
    sold_net_prices = net_price.loc[is_sold_window].dropna()

    active_min = float(active_prices.min()) if not active_prices.empty else np.nan
    active_max = float(active_prices.max()) if not active_prices.empty else np.nan
    pending_min = float(pending_prices.min()) if not pending_prices.empty else np.nan
    pending_max = float(pending_prices.max()) if not pending_prices.empty else np.nan
    sold_net_min = float(sold_net_prices.min()) if not sold_net_prices.empty else np.nan
    sold_net_max = float(sold_net_prices.max()) if not sold_net_prices.empty else np.nan

    dom_series = df.loc[is_sold_window, "_dom"].dropna()
    avg_dom = float(dom_series.mean()) if not dom_series.empty else np.nan
    median_dom = float(dom_series.median()) if not dom_series.empty else np.nan

    doc = Document()
    title = doc.add_paragraph("Momentum Report")
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pdate = doc.add_paragraph(f"Report Date: {today.strftime('%B %d, %Y')}")
    pdate.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph("Summary", style=None).runs[0].bold = True
    doc.add_paragraph(f"Active: {active_count}")
    doc.add_paragraph(f"Pending: {pending_count}")
    doc.add_paragraph(f"Solds (last {days}d): {sold_window_count}")
    doc.add_paragraph(f"Months of Inventory (MOI): {moi:.3f}" if not np.isnan(moi) else "Months of Inventory (MOI): N/A")

    doc.add_paragraph("Price Ranges", style=None).runs[0].bold = True
    doc.add_paragraph(
        f"Active List Price Range: ${active_min:,.0f} – ${active_max:,.0f}"
        if not (np.isnan(active_min) or np.isnan(active_max)) else "Active List Price Range: N/A"
    )
    doc.add_paragraph(
        f"Pending List Price Range: ${pending_min:,.0f} – ${pending_max:,.0f}"
        if not (np.isnan(pending_min) or np.isnan(pending_max)) else "Pending List Price Range: N/A"
    )
    doc.add_paragraph(
        f"Closed Net Price Range ({days}d): ${sold_net_min:,.0f} – ${sold_net_max:,.0f}"
        if not (np.isnan(sold_net_min) or np.isnan(sold_net_max)) else f"Closed Net Price Range ({days}d): N/A"
    )

    doc.add_paragraph("Days in MLS (Closed, Window)", style=None).runs[0].bold = True
    doc.add_paragraph(f"Average DaysInMLS: {avg_dom:.1f}" if not np.isnan(avg_dom) else "Average DaysInMLS: N/A")
    doc.add_paragraph(f"Median DaysInMLS: {median_dom:.1f}" if not np.isnan(median_dom) else "Median DaysInMLS: N/A")

    doc.add_paragraph("")
    doc.add_paragraph("Formula", style=None).runs[0].bold = True
    doc.add_paragraph(f"MOI = Active / ((Solds_{days}d / 3) + Pending)")

    doc.add_paragraph("")
    doc.add_paragraph("Disclaimer", style=None).runs[0].bold = True
    doc.add_paragraph(
        "This report is generated from the provided MLS export. Field names are expected to match the fixed set; "
        "optional fields like Seller Concessions and DaysInMLS may be omitted."
    )

    file_bytes = io.BytesIO()
    doc.save(file_bytes)
    file_bytes.seek(0)
    return file_bytes

if uploaded is not None:
    # Basic validation for required columns
    try:
        df = pd.read_csv(uploaded, dtype=str)
    except Exception as e:
        st.error(f"Could not read CSV: {e}")
        st.stop()

    required = ["Mls Status", "Close Date", "List Price", "Close Price"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        st.error(f"Missing required columns: {missing}")
        st.stop()

    st.success("CSV loaded. Ready to generate report.")
    if st.button("Generate DOCX Report"):
        try:
            buf = build_docx(df, int(days_window))
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Momentum_Report_{stamp}.docx"
            st.download_button("Download Report", data=buf, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.info("Report generated successfully.")
        except Exception as e:
            st.error(f"Failed to generate report: {e}")
else:
    st.info("Upload a CSV to begin.")
